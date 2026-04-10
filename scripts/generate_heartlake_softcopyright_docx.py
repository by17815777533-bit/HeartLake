from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION_START
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt


ROOT = Path(__file__).resolve().parents[1]
SCREENSHOT_DIR = ROOT / "output" / "playwright"
DEFAULT_OUTPUT = Path(
    "/home/bai/下载/0、软著申请模板 (1)/0、软著申请模板/HeartLake_软件说明书_软著版.docx"
)


SECTIONS = [
    {
        "title": "一、系统概述",
        "paragraphs": [
            "心湖匿名心理关怀系统软件 V1.0 是一款面向移动端用户的匿名情绪表达与陪伴类软件。用户无需提供真实姓名、手机号等敏感身份信息，即可通过匿名登录进入系统，在相对低压力、低暴露的环境中记录心情、浏览内容、接收回应，并逐步建立温和的互动关系。",
            "软件围绕“投石表达、涟漪共鸣、纸船回应、关系建立、情绪感知、心理关怀、安全治理”几条主线组织功能，既强调匿名使用体验，也兼顾账号恢复、隐私设置、安全港引导、情绪趋势查看与 AI 陪伴等能力，适合用于日常情绪记录、匿名倾诉和基础心理支持场景。",
            "本说明书按照用户实际使用顺序编写，从软件启动、匿名登录开始，依次说明观湖浏览、通知查看、内容发现、好友关系、个人中心、情绪日历、守护者、安全港、心理咨询、湖神聊天、隐私设置以及退出使用等操作流程。",
        ],
    },
    {
        "title": "二、运行环境与使用对象",
        "paragraphs": [
            "本软件前端采用 Flutter 构建，适用于安卓移动端，也支持当前版本的 Web 运行形态用于联调与演示。系统后端提供匿名认证、内容交互、推荐、情绪分析、隐私设置、咨询服务与实时消息等接口支持。用户仅需具备可联网的终端设备即可完成大部分功能操作。",
            "普通用户进入软件后，可匿名生成账号并保存恢复密钥；在后续使用过程中，可浏览湖面内容、查看推荐、发送纸船、查看情绪趋势、设置守护者、使用安全港与湖神陪伴功能。由于系统采用匿名身份机制，因此首次进入后建议优先妥善保存恢复密钥，以便后续找回已有账户。",
        ],
    },
    {
        "title": "三、软件启动与匿名登录",
        "paragraphs": [
            "用户首次打开软件后，会进入登录页。系统提供“匿名进入心湖”和“恢复我的账号”两种入口。其中，匿名进入用于创建新的匿名账户，恢复账号用于通过已经保存的恢复密钥找回旧账户。",
            "选择匿名进入后，系统会自动生成当前匿名身份，并返回用户昵称、访问令牌、刷新令牌、会话编号以及恢复密钥。恢复密钥是账户找回的唯一凭证，建议用户立即复制并妥善保管。只有在确认已经保存恢复密钥后，才继续进入后续页面。",
        ],
        "image": "auth-page.png",
        "caption": "图 1 软件匿名登录页",
    },
    {
        "title": "四、进入观湖首页",
        "paragraphs": [
            "登录完成后，用户进入软件首页“观湖”页面。该页面用于展示湖面中的公开内容，也就是其他用户投递的“石头”。每条石头包含用户昵称、情绪标签、正文内容以及互动入口，用户可以在此持续浏览、下拉刷新并了解当前社区的表达状态。",
            "首页顶部提供湖神入口、共鸣发现入口与通知入口；底部导航栏则提供“观湖、共鸣、投石、好友、倒影”等主功能切换。对于普通用户而言，观湖页是进入系统后的主浏览界面，也是后续互动与发现内容的基础入口。",
        ],
        "image": "home.png",
        "caption": "图 2 软件观湖首页",
    },
    {
        "title": "五、通知中心",
        "paragraphs": [
            "当用户收到纸船、涟漪、关系变化或系统消息时，可以通过通知入口进入通知中心查看。通知页会集中呈现近期与本人有关的动态，帮助用户快速了解自己的内容是否获得反馈，以及当前是否存在需要关注的提醒事项。",
            "在当前截图示例中，通知页处于空状态，界面会以“暂无通知，心湖很安静”的形式提示用户当前没有新的消息；实际使用过程中，当存在互动或系统推送时，该区域会展示相应记录。",
        ],
        "image": "notifications.png",
        "caption": "图 3 通知中心页面",
    },
    {
        "title": "六、共鸣发现",
        "paragraphs": [
            "点击底部导航中的“共鸣”或从首页进入发现页后，用户可以查看系统整理的热门内容、关键词查找入口以及湖神陪伴分区。该页面的目标是帮助用户从“我正在浏览什么”扩展到“我想继续寻找什么内容”，形成更主动的内容发现流程。",
            "发现页包含热门、找心声、湖神陪伴等标签页。用户可以按分区浏览内容，也可以在内容不足时看到明确的空状态提示，不会出现静默失败的情况。对于新用户而言，这里是快速熟悉社区氛围和寻找相似情绪表达的重要入口。",
        ],
        "image": "discover.png",
        "caption": "图 4 共鸣发现页面",
    },
    {
        "title": "七、个性化推荐",
        "paragraphs": [
            "在“为你而来”页面中，系统会根据用户的情绪特征、历史互动和内容匹配结果，生成更贴近当前状态的推荐列表。页面中的推荐卡片会展示石头内容、情绪匹配程度以及基础互动数据，帮助用户快速找到更有共鸣的内容。",
            "个性化推荐功能的意义在于减少用户在海量内容中的无效浏览，让用户更容易遇到和自己当前情绪接近的表达，从而获得被理解、被看见或被陪伴的体验。这类推荐结果既可以作为阅读入口，也能进一步引导用户进入详情页或发起互动。",
        ],
        "image": "personalized.png",
        "caption": "图 5 个性化推荐页面",
    },
    {
        "title": "八、个人中心与资料查看",
        "paragraphs": [
            "用户可从底部导航进入“倒影”页面，也就是个人中心。该页面汇总展示当前账号昵称、头像入口、投石数量、收到纸船数量、发送纸船数量和相伴天数等信息，同时提供情绪日历、情绪趋势、热力图、我的涟漪、守护者、安全港、心理咨询、隐私设置、帮助中心等统一入口。",
            "个人中心的价值在于把与“我”有关的内容、状态、统计和工具汇总到一个位置，方便用户快速回看自己的使用轨迹，也方便后续进行更细致的情绪查看和隐私管理。",
        ],
        "image": "profile.png",
        "caption": "图 6 个人中心页面",
    },
    {
        "title": "九、好友与临时好友",
        "paragraphs": [
            "软件支持好友关系与临时好友机制。用户进入好友页后，可以查看当前已经建立的正式好友关系，也可以通过“临时好友”入口进入短时关系场景。临时好友通常具有 24 小时有效期，更适合用于从一次纸船互动过渡到更进一步的连接。",
            "在当前示例中，由于账号刚刚生成，因此好友列表为空，页面会给出明确的空状态提示。实际使用中，用户可通过互动逐步建立临时好友关系，并在持续交流后升级为更稳定的好友关系。",
        ],
        "image": "friends.png",
        "caption": "图 7 好友页面",
    },
    {
        "title": "十、我的石头与内容管理",
        "paragraphs": [
            "用户可以在个人中心进一步查看“我的石头”，也就是自己已经发布过的内容。该页面用于回看个人表达记录，方便用户检查过往写下的情绪内容是否已经成功发布，以及后续是否收到了互动反馈。",
            "在当前示例中，用户尚未投递新的石头，因此页面展示空状态。实际使用中，该区域会列出自己已发布的石头内容，并随着纸船、涟漪等实时变化同步更新互动结果。",
        ],
        "image": "my-stones.png",
        "caption": "图 8 我的石头页面",
    },
    {
        "title": "十一、情绪日历",
        "paragraphs": [
            "情绪日历页面用于按日期维度回看用户每日情绪变化情况。页面顶部会汇总当前月份中的主要情绪分布，日历区域则用于按日展示情绪记录，让用户更容易识别近期心情是否稳定，是否存在明显波动或持续低落的趋势。",
            "通过该功能，用户可以从零散的即时感受，过渡到更结构化的长期观察。对于需要自我觉察或持续记录情绪的人来说，情绪日历是非常直观的工具页面。",
        ],
        "image": "emotion-calendar.png",
        "caption": "图 9 情绪日历页面",
    },
    {
        "title": "十二、情绪趋势",
        "paragraphs": [
            "情绪趋势页面主要用于展示用户近期情绪变化方向和总体倾向。系统会综合已有样本进行趋势判断，并在样本不足时明确提示用户继续记录情绪，以便后续形成更准确的分析结果。同时，页面还会显示差分隐私保护状态，让用户知道统计结果是在隐私保护前提下生成的。",
            "相较于情绪日历偏重单日记录，情绪趋势更强调阶段性变化判断，能够帮助用户回答“我最近整体变好了还是变差了”这样的问题，因此更适合做一段时间后的整体回顾。",
        ],
        "image": "emotion-trends.png",
        "caption": "图 10 情绪趋势页面",
    },
    {
        "title": "十三、守护者",
        "paragraphs": [
            "守护者页面用于承接软件中的“守护关系”能力。用户可以在此查看自己是否已经具备守护者资格、当前共鸣点、优质涟漪数、温暖纸船数以及情感洞察结果。系统还会在这里给出近期情绪概况和简短建议，帮助用户理解自己当前的状态。",
            "当用户逐步累积互动与共鸣之后，可以更稳定地承担守护者角色，为他人提供温和支持；反过来，用户自己也可以通过此页理解守护机制的意义，并进入与湖神对话等后续功能。",
        ],
        "image": "guardian.png",
        "caption": "图 11 守护者页面",
    },
    {
        "title": "十四、安全港湾",
        "paragraphs": [
            "安全港湾是软件中的重点支持页面，用于在用户需要更直接帮助时提供稳定、可信的求助入口。页面中包含全国心理援助热线、专业心理机构干预电话、生命热线等资源，同时配有深呼吸练习、五感着陆、情绪日记等自助工具。",
            "该功能面向的是“需要更立即支持”的场景。用户在这里既可以联系专业资源，也可以先使用轻量自助工具平复当下情绪。页面底部会强调匿名与记录保护，降低用户在求助场景下的顾虑。",
        ],
        "image": "safe-harbor.png",
        "caption": "图 12 安全港湾页面",
    },
    {
        "title": "十五、心理咨询",
        "paragraphs": [
            "心理咨询页面展示了软件中的咨询预约与会话预览能力。当前版本中的页面重点在于演示预约体验和版面结构，用户可以先选择陪伴方向、预约时间草稿以及想聊的话题，以熟悉正式咨询流程的交互方式。",
            "此页面适合在用户已经完成基础表达和浏览之后，继续向更专业支持过渡的场景。通过相对温和的预约草稿与方向选择设计，系统降低了用户直接进入正式咨询流程的心理门槛。",
        ],
        "image": "consultation.png",
        "caption": "图 13 心理咨询页面",
    },
    {
        "title": "十六、湖神聊天",
        "paragraphs": [
            "湖神聊天页面是软件中的 AI 陪伴入口。用户可以把想说的话发送给湖神，由系统结合情绪分析、内容审核和历史会话能力给出回应；在风险较高或内容不适合继续对话时，系统也会配合安全港机制做进一步引导。",
            "在当前运行示例中，页面给出了历史消息加载失败的显式提示和重试入口。这种设计可以避免系统异常时向用户伪装成功，能够更清楚地提示用户当前状态，并保护后续使用体验。",
        ],
        "image": "lake-god-chat.png",
        "caption": "图 14 湖神聊天页面",
    },
    {
        "title": "十七、隐私与安全设置",
        "paragraphs": [
            "隐私与安全页面用于管理软件中的关键隐私开关。用户可以控制是否显示在线状态、是否允许陌生人发送纸船、是否对陌生人公开自己的资料；同时，还可以在此申请导出个人数据、停用账号或执行永久删除等操作。",
            "由于心湖强调匿名与隐私保护，因此这一页面是用户长期使用过程中的重要设置页。用户可以根据自己的舒适边界调节可见性与互动范围，让软件更贴合个人需求。",
        ],
        "image": "privacy-settings.png",
        "caption": "图 15 隐私与安全设置页面",
    },
    {
        "title": "十八、投石发布说明",
        "paragraphs": [
            "在主界面底部导航中，点击“投石”即可进入内容发布流程。用户需要先选择当前心情，再输入想表达的文字内容，系统会结合本地情绪识别与内容审核能力给出参考提示，帮助用户更清晰地表达情绪，也避免明显不当内容直接发布。",
            "当用户确认发布后，石头会被投入湖中，并进入首页内容流与自己的内容记录页。若发布成功，系统会提示“石头已投入湖中，等待涟漪”，这意味着其他用户后续可以对该内容发起涟漪或纸船互动。",
        ],
        "placeholder": "截图占位：投石发布页。建议补充“选择心情 + 输入内容 + 发布成功提示”界面截图。",
    },
    {
        "title": "十九、涟漪、纸船与详情互动说明",
        "paragraphs": [
            "当用户浏览到某条石头后，可以进入详情页查看完整内容，并进行两类核心互动：一类是“涟漪”，用于表示共鸣或认可；另一类是“纸船”，用于发送更温和、具有回应意味的匿名留言。系统会在服务端确认后同步更新互动状态，避免本地显示与实际结果不一致。",
            "对于内容作者而言，收到的纸船和涟漪会回流到通知中心、我的石头、收到的纸船等页面；对于互动者而言，纸船互动还有机会进一步形成临时好友关系，为后续持续交流创造条件。",
        ],
        "placeholder": "截图占位：石头详情页或纸船互动页。建议补充“涟漪按钮、纸船输入框、纸船列表”实际界面截图。",
    },
    {
        "title": "二十、结束使用与账号恢复提醒",
        "paragraphs": [
            "当用户完成本次使用后，可直接退出当前页面或在个人中心执行登出操作。由于软件采用匿名账号体系，因此即便用户不提供实名信息，也仍然需要依靠恢复密钥来保证账户可持续使用。若后续更换设备或本地数据丢失，用户可通过恢复密钥找回原有账户。",
            "建议用户在结束使用前再次确认恢复密钥是否已经妥善保存，并根据自身需要检查隐私设置是否处于合适状态。对于涉及长期陪伴、内容积累和情绪记录的用户来说，保存好恢复密钥与合理设置隐私范围同样重要。",
        ],
    },
]


def set_chinese_font(run, font_name: str, size: float | None = None, bold: bool | None = None) -> None:
    run.font.name = font_name
    run._element.rPr.rFonts.set(qn("w:eastAsia"), font_name)
    if size is not None:
        run.font.size = Pt(size)
    if bold is not None:
        run.bold = bold


def apply_doc_style(doc: Document) -> None:
    section = doc.sections[0]
    section.page_width = Cm(21)
    section.page_height = Cm(29.7)
    section.top_margin = Cm(2.54)
    section.bottom_margin = Cm(2.54)
    section.left_margin = Cm(2.54)
    section.right_margin = Cm(2.54)

    normal = doc.styles["Normal"]
    normal.font.name = "宋体"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
    normal.font.size = Pt(12)
    normal.paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
    normal.paragraph_format.space_after = Pt(6)

    for style_name, font_name, size in [
        ("Title", "黑体", 22),
        ("Heading 1", "黑体", 16),
        ("Heading 2", "黑体", 14),
    ]:
        style = doc.styles[style_name]
        style.font.name = font_name
        style._element.rPr.rFonts.set(qn("w:eastAsia"), font_name)
        style.font.size = Pt(size)


def add_cover(doc: Document) -> None:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.space_before = Pt(80)
    run = p.add_run("心湖匿名心理关怀系统软件 V1.0")
    set_chinese_font(run, "黑体", 22, True)

    p2 = doc.add_paragraph()
    p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p2.space_before = Pt(18)
    run = p2.add_run("软件说明书")
    set_chinese_font(run, "黑体", 24, True)

    p3 = doc.add_paragraph()
    p3.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p3.space_before = Pt(18)
    run = p3.add_run("（软著申报图文说明初稿）")
    set_chinese_font(run, "宋体", 14)

    spacer = doc.add_paragraph()
    spacer.space_before = Pt(140)

    p4 = doc.add_paragraph()
    p4.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p4.add_run("软件名称：心湖匿名心理关怀系统软件")
    set_chinese_font(run, "宋体", 14)

    p5 = doc.add_paragraph()
    p5.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p5.add_run("版本号：V1.0")
    set_chinese_font(run, "宋体", 14)

    p6 = doc.add_paragraph()
    p6.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p6.add_run("编制日期：2026 年 4 月")
    set_chinese_font(run, "宋体", 14)

    doc.add_page_break()


def add_directory(doc: Document) -> None:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("目录")
    set_chinese_font(run, "黑体", 18, True)

    for item in SECTIONS:
        p = doc.add_paragraph(style="Normal")
        p.paragraph_format.left_indent = Cm(0.8)
        run = p.add_run(item["title"])
        set_chinese_font(run, "宋体", 12)

    note = doc.add_paragraph()
    note.paragraph_format.space_before = Pt(10)
    run = note.add_run("说明：本说明书按用户实际使用流程编写，文中截图来自当前版本运行界面；个别尚未补齐的页面已预留截图占位。")
    set_chinese_font(run, "宋体", 11)
    doc.add_page_break()


def add_heading(doc: Document, text: str) -> None:
    p = doc.add_paragraph()
    p.style = "Heading 1"
    run = p.add_run(text)
    set_chinese_font(run, "黑体", 16, True)


def add_body_paragraph(doc: Document, text: str) -> None:
    p = doc.add_paragraph(style="Normal")
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    run = p.add_run(text)
    set_chinese_font(run, "宋体", 12)
    p.paragraph_format.first_line_indent = Cm(0.84)


def add_caption(doc: Document, text: str) -> None:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(text)
    set_chinese_font(run, "宋体", 10.5)


def shade_cell(cell, fill: str = "F7F7F7") -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def add_placeholder(doc: Document, text: str) -> None:
    table = doc.add_table(rows=1, cols=1)
    table.style = "Table Grid"
    cell = table.cell(0, 0)
    shade_cell(cell)
    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(18)
    p.paragraph_format.space_after = Pt(18)
    run = p.add_run(text)
    set_chinese_font(run, "宋体", 11)
    doc.add_paragraph()


def add_image_or_placeholder(doc: Document, image_name: str | None, caption: str | None, placeholder: str | None) -> None:
    if image_name:
        image_path = SCREENSHOT_DIR / image_name
        if image_path.exists():
            doc.add_picture(str(image_path), width=Cm(15.5))
            if caption:
                add_caption(doc, caption)
            return
    if placeholder:
        add_placeholder(doc, placeholder)


def build_doc(output_path: Path) -> None:
    doc = Document()
    apply_doc_style(doc)
    add_cover(doc)
    add_directory(doc)

    for index, section in enumerate(SECTIONS):
        if index > 0:
            doc.add_page_break()
        add_heading(doc, section["title"])
        for paragraph in section["paragraphs"]:
            add_body_paragraph(doc, paragraph)
        add_image_or_placeholder(
            doc,
            section.get("image"),
            section.get("caption"),
            section.get("placeholder"),
        )

    output_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(output_path))


if __name__ == "__main__":
    build_doc(DEFAULT_OUTPUT)
